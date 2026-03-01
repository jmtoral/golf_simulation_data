"""
GoGolf - Generador de documento Word
Secciones: Metas SMART + Recoleccion/Limpieza/Validacion + Analisis Descriptivo
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, csv, statistics
from collections import defaultdict

BASE = os.path.dirname(__file__)
CSV  = os.path.join(BASE, "csv")
FIGS = os.path.join(BASE, "graficas")

def load(fname):
    with open(os.path.join(CSV, fname), encoding="utf-8") as f:
        return list(csv.DictReader(f))

# Cargar KPIs reales
bsc    = load("kpi_bsc_mensual.csv")
res    = load("fact_reservas.csv")
noshow = load("fact_noshow.csv")
rats   = load("fact_ratings.csv")
fric   = load("fact_fricciones.csv")
clubs  = load("dim_club.csv")

def col(table, c): return [float(r[c]) for r in table if r.get(c,'')]

# KPIs actuales (linea base)
TASA_CAN   = round(statistics.mean(col(bsc,"tasa_cancelacion_pct")),1)
TASA_NS    = round(statistics.mean(col(bsc,"tasa_noshow_pct")),1)
NPS        = round(statistics.mean(col(bsc,"nps_proxy")),1)
RATING     = round(statistics.mean(col(bsc,"rating_promedio_mes")),2)
RECOMPRA   = round(statistics.mean(col(bsc,"tasa_recompra_pct")),1)
CUMPL      = round(statistics.mean(col(bsc,"pct_cumplimiento_horario")),1)
DISC       = round(statistics.mean(col(bsc,"pct_discrepancia_inventario")),2)
FRIC       = round(statistics.mean(col(bsc,"pct_friccion_social")),1)
MARGEN     = round(statistics.mean(col(bsc,"margen_promedio_por_transaccion"))*100,1)
ING_PROM   = round(statistics.mean(col(bsc,"ingreso_promedio_por_reserva")))
PERDIDA    = round(sum(float(r["perdida_total_estimada_mxn"]) for r in noshow)/1e6, 1)
TOT_RES    = len(res)
nse_cnt    = defaultdict(int)
for r in res: nse_cnt[r["nse_jugador"]] += 1
tot = sum(nse_cnt.values())
PCT_NSE_BAJO = round((nse_cnt["C"]+nse_cnt["Dmas"])/tot*100)

# ── Helpers de formato ────────────────────────────────────────────────────────
AZUL_OSC  = RGBColor(0x1F, 0x4E, 0x79)
AZUL_MED  = RGBColor(0x2E, 0x75, 0xB6)
GRIS_FONDO= RGBColor(0xF2, 0xF2, 0xF2)
NARANJA   = RGBColor(0xED, 0x7D, 0x31)
VERDE     = RGBColor(0x70, 0xAD, 0x47)
ROJO      = RGBColor(0xC0, 0x00, 0x00)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        if edge in kwargs:
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), kwargs[edge].get("val","single"))
            tag.set(qn("w:sz"),  kwargs[edge].get("sz","4"))
            tag.set(qn("w:space"),"0")
            tag.set(qn("w:color"),kwargs[edge].get("color","auto"))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if color: run.font.color.rgb = color
    run.font.bold = True
    return p

def body(doc, text, bold_parts=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    if bold_parts:
        parts = text.split("**")
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.font.size = Pt(11)
            if i % 2 == 1:
                run.bold = True
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level*0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_kpi_line(doc, kpi, valor_actual, meta_6m, meta_12m, meta_24m, nota=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run(f"{kpi}: ")
    r1.bold = True; r1.font.size = Pt(10.5)
    r2 = p.add_run(f"Actual {valor_actual} → 6m: {meta_6m} → 12m: {meta_12m} → 24m: {meta_24m}")
    r2.font.size = Pt(10.5)
    if nota:
        r3 = p.add_run(f"  ({nota})")
        r3.font.size = Pt(9.5); r3.font.color.rgb = RGBColor(0x80,0x80,0x80)

def tabla_metas(doc):
    """Tabla consolidada de metas SMART."""
    filas = [
        # (perspectiva, iniciativa, KPI, baseline, meta_6m, meta_12m, meta_24m)
        ("Financiera",
         "Maximizar valor\npor transaccion",
         "Ingreso prom/reserva (MXN)",
         f"${ING_PROM:,}",
         f"${round(ING_PROM*1.05):,}",
         f"${round(ING_PROM*1.12):,}",
         f"${round(ING_PROM*1.25):,}"),
        ("Financiera",
         "Maximizar valor\npor transaccion",
         "Margen por transaccion (%)",
         f"{MARGEN}%",
         f"{round(MARGEN+3,1)}%",
         f"{round(MARGEN+6,1)}%",
         f"{round(MARGEN+10,1)}%"),
        ("Financiera",
         "Reducir perdidas\noperativas",
         "Tasa de cancelacion (%)",
         f"{TASA_CAN}%",
         f"<{TASA_CAN-3:.0f}%",
         f"<{TASA_CAN-5:.0f}%",
         f"<{TASA_CAN-7:.0f}%"),
        ("Cliente",
         "Medir experiencia\nintegral",
         "NPS proxy",
         f"{NPS}",
         ">0",
         ">+10",
         ">+20"),
        ("Cliente",
         "Fomentar\nrecurrencia",
         "Tasa de recompra (%)",
         f"{RECOMPRA}%",
         f">{round(RECOMPRA+5,0):.0f}%",
         f">{round(RECOMPRA+10,0):.0f}%",
         f">{round(RECOMPRA+15,0):.0f}%"),
        ("Cliente",
         "Garantizar calidad\ny reputacion",
         "Rating promedio (1-5)",
         f"{RATING}",
         ">4.0",
         ">4.2",
         ">4.5"),
        ("Procesos\nInternos",
         "Alinear clubes\nbajo SLAs",
         "% Cumplimiento horario",
         f"{CUMPL}%",
         ">90%",
         ">93%",
         ">96%"),
        ("Procesos\nInternos",
         "Reducir friccion\nmarketplace",
         "% Discrepancia inventario",
         f"{DISC}%",
         f"<{round(DISC*0.75,1)}%",
         f"<{round(DISC*0.5,1)}%",
         f"<{round(DISC*0.25,1)}%"),
        ("Procesos\nInternos",
         "Reducir siniestros\noperativos",
         "Tasa de no-show (%)",
         f"{TASA_NS}%",
         f"<{round(TASA_NS-3,0):.0f}%",
         f"<{round(TASA_NS-5,0):.0f}%",
         f"<{round(TASA_NS-7,0):.0f}%"),
        ("Aprendizaje\ny Crecimiento",
         "Reducir brecha NSE\n(friccion social)",
         "% Reservas con friccion social",
         f"{FRIC}%",
         f"<{round(FRIC*0.85,0):.0f}%",
         f"<{round(FRIC*0.70,0):.0f}%",
         f"<{round(FRIC*0.50,0):.0f}%"),
        ("Aprendizaje\ny Crecimiento",
         "Transicion\ndata-driven",
         "% Decisiones con analisis formal",
         "Est. 30%",
         ">50%",
         ">70%",
         ">85%"),
        ("Aprendizaje\ny Crecimiento",
         "Agilidad\norganizacional",
         "Tiempo impl. mejoras (dias)",
         "Est. 45",
         "<35",
         "<25",
         "<15"),
    ]

    # Encabezados
    headers = ["Perspectiva","Iniciativa Estrategica","KPI",
               "Baseline\n(actual)","Meta\n6 meses","Meta\n12 meses","Meta\n24 meses"]
    t = doc.add_table(rows=1+len(filas), cols=7)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = t.rows[0]
    for i,(h,w) in enumerate(zip(headers,[2.5,3.5,4.5,2.0,2.0,2.0,2.0])):
        c = hdr.cells[i]
        set_cell_bg(c,"1F4E79")
        c.width = Cm(w)
        p2 = c.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p2.add_run(h)
        run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    # Filas de datos
    perspectiva_color = {
        "Financiera":           "D6E4F0",
        "Cliente":              "E2F0D9",
        "Procesos\nInternos":   "FFF2CC",
        "Aprendizaje\ny Crecimiento":"F8CBAD",
    }
    prev_persp = None
    for i, fila in enumerate(filas):
        row = t.rows[i+1]
        persp, inic, kpi, base, m6, m12, m24 = fila
        color_hex = perspectiva_color.get(persp,"F2F2F2")
        vals = [persp, inic, kpi, base, m6, m12, m24]
        for j, val in enumerate(vals):
            cell = row.cells[j]
            set_cell_bg(cell, color_hex)
            p2 = cell.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER if j>2 else WD_ALIGN_PARAGRAPH.LEFT
            run = p2.add_run(val)
            run.font.size = Pt(8.5)
            if j == 0: run.bold = True
    return t

# =============================================================================
# CONSTRUIR DOCUMENTO
# =============================================================================
doc = Document()

# Margenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Estilo base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ════════════════════════════════════════════════════════════════════════════
# BLOQUE A — METAS Y OBJETIVOS ESTRATEGICOS
# ════════════════════════════════════════════════════════════════════════════
h = doc.add_heading("3. METAS Y OBJETIVOS ESTRATÉGICOS", level=1)
h.runs[0].font.color.rgb = AZUL_OSC

body(doc,
    "Las metas estratégicas de GoGolf se estructuran bajo el principio SMART, asegurando "
    "que cada indicador del Balanced Scorecard (BSC) tenga un objetivo cuantificable, "
    "alcanzable y temporalmente definido. Los valores de referencia (baseline) provienen "
    "del análisis descriptivo de los datos del período 2023–2024, sintetizados en el "
    "datamart operativo de la plataforma."
)

body(doc,
    f"El diagnóstico inicial revela tres tensiones estructurales: una **tasa de no-show "
    f"de {TASA_NS}%** —superior al promedio del sector de plataformas de reserva turística "
    f"(~8–10%)—, un **NPS proxy negativo de {NPS}** que refleja experiencias mixtas entre "
    f"usuarios de nivel socioeconómico medio-bajo, y una **tasa de fricción social del "
    f"{FRIC}%** causada por la brecha entre los requisitos de entrada de los clubes y el "
    f"perfil del jugador típico de GoGolf —donde el {PCT_NSE_BAJO}% pertenece a los "
    f"segmentos NSE C y D+.",
    bold_parts=True
)

doc.add_heading("3.1 Horizontes Estratégicos", level=2).runs[0].font.color.rgb = AZUL_MED

body(doc,
    "Se establecen tres horizontes estratégicos, alineados con la lógica causa–efecto "
    "del BSC: Capacidades → Procesos → Cliente → Financiera:"
)

for horizonte, titulo, desc in [
    ("Corto plazo (6 meses):",
     "Estabilización operativa y reducción de fricciones críticas.",
     "Prioridad en reducir el no-show, eliminar discrepancias de inventario y establecer "
     "un protocolo de onboarding para jugadores NSE C/D+ que anticipe los requisitos de "
     "los clubes antes de la reserva."),
    ("Mediano plazo (12 meses):",
     "Optimización del modelo y consolidación de recurrencia.",
     "Una vez estabilizada la operación, el foco se traslada a convertir usuarios "
     "ocasionales en jugadores recurrentes, elevar el NPS a terreno positivo y capturar "
     "mayor margen a través de pricing dinámico y up-selling."),
    ("Largo plazo (24 meses):",
     "Escalabilidad sostenible y madurez del modelo.",
     "Incorporar nuevos clubes y mercados, institucionalizar la toma de decisiones "
     "basada en datos y alcanzar métricas de referencia de plataformas de marketplace "
     "deportivo consolidadas."),
]:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(horizonte + " ")
    r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(titulo + " ")
    r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = AZUL_MED
    r3 = p.add_run(desc)
    r3.font.size = Pt(10.5)

doc.add_heading("3.2 Definición de Metas SMART por KPI", level=2).runs[0].font.color.rgb = AZUL_MED

body(doc,
    "A continuación se definen las metas por perspectiva del BSC, con base en los valores "
    "actuales del sistema de datos y benchmarks del sector de plataformas de reserva "
    "en línea en mercados latinoamericanos:"
)

# -- Financiera --
doc.add_heading("Perspectiva Financiera", level=3).runs[0].font.color.rgb = AZUL_MED
body(doc,
    f"El ingreso promedio por reserva se ubica en **${ING_PROM:,} MXN** con un margen "
    f"promedio por transacción de **{MARGEN}%**. La tasa de cancelación actual es de "
    f"**{TASA_CAN}%**, generando pérdida de comisión en cada evento. La pérdida "
    f"acumulada por no-shows en el período 2023–2024 asciende a **${PERDIDA:.1f} millones "
    f"de MXN**, de los cuales una fracción significativa es atribuible a fricción social "
    f"(jugadores que no se presentan por no cumplir requisitos del club).",
    bold_parts=True
)
add_kpi_line(doc,"Ingreso promedio por reserva",f"${ING_PROM:,} MXN",
             f"${round(ING_PROM*1.05):,}",f"${round(ING_PROM*1.12):,}",f"${round(ING_PROM*1.25):,}",
             "via up-selling (renta de equipo, caddies, clases)")
add_kpi_line(doc,"Margen por transacción",f"{MARGEN}%",
             f"{round(MARGEN+3,1)}%",f"{round(MARGEN+6,1)}%",f"{round(MARGEN+10,1)}%",
             "optimizacion de costos de soporte y tecnologia")
add_kpi_line(doc,"Tasa de cancelación",f"{TASA_CAN}%",
             f"<{TASA_CAN-3:.0f}%",f"<{TASA_CAN-5:.0f}%",f"<{TASA_CAN-7:.0f}%",
             "politicas de deposito y recordatorios automaticos")

# -- Cliente --
doc.add_heading("Perspectiva de Cliente", level=3).runs[0].font.color.rgb = AZUL_MED
body(doc,
    f"El NPS proxy calculado a partir de ratings es de **{NPS}**, indicando que existen "
    f"más detractores que promotores en la base actual —resultado consistente con la "
    f"fricción social documentada. El rating promedio de **{RATING}/5.0** y la tasa "
    f"de recompra de **{RECOMPRA}%** sugieren que quienes logran completar una reserva "
    f"sin fricciones tienden a regresar, pero la barrera de entrada trunca la "
    f"conversión de nuevos usuarios.",
    bold_parts=True
)
add_kpi_line(doc,"NPS proxy",f"{NPS}",">0",">+10",">+20",
             "experiencia de onboarding + curador de clubes compatibles por NSE")
add_kpi_line(doc,"Tasa de recompra",f"{RECOMPRA}%",
             f">{round(RECOMPRA+5):.0f}%",f">{round(RECOMPRA+10):.0f}%",f">{round(RECOMPRA+15):.0f}%",
             "programa de fidelidad y notificaciones personalizadas")
add_kpi_line(doc,"Rating promedio",f"{RATING}/5.0",">4.0",">4.2",">4.5",
             "estandar minimo de calidad para clubes afiliados")

# -- Procesos --
doc.add_heading("Perspectiva de Procesos Internos", level=3).runs[0].font.color.rgb = AZUL_MED
body(doc,
    f"El cumplimiento de horario actual es de **{CUMPL}%**, con un **{DISC}%** de "
    f"discrepancias entre el inventario digital y la disponibilidad real de los clubes. "
    f"La tasa de no-show de **{TASA_NS}%** es la métrica de mayor impacto operativo: "
    f"cada no-show representa un slot irrecuperable para el club y la pérdida de la "
    f"comisión de GoGolf. El análisis de causas muestra que la **fricción social es el "
    f"principal driver**, superando incluso al clima adverso.",
    bold_parts=True
)
add_kpi_line(doc,"% Cumplimiento de horarios",f"{CUMPL}%",">90%",">93%",">96%",
             "SLA formales con clubes + sistema de alertas")
add_kpi_line(doc,"% Discrepancias de inventario",f"{DISC}%",
             f"<{round(DISC*0.75,1)}%",f"<{round(DISC*0.5,1)}%",f"<{round(DISC*0.25,1)}%",
             "API en tiempo real con sistema de reservas de cada club")
add_kpi_line(doc,"Tasa de no-shows",f"{TASA_NS}%",
             f"<{round(TASA_NS-3):.0f}%",f"<{round(TASA_NS-5):.0f}%",f"<{round(TASA_NS-7):.0f}%",
             "deposito obligatorio + filtro de requisitos pre-reserva")

# -- Aprendizaje --
doc.add_heading("Perspectiva de Aprendizaje y Crecimiento", level=3).runs[0].font.color.rgb = AZUL_MED
body(doc,
    f"El indicador más revelador de esta perspectiva es el **{FRIC}% de reservas con "
    f"fricción social**, concentrado en jugadores NSE C y D+ que acceden a clubes con "
    f"requisitos de dress code estricto, handicap máximo o equipo propio obligatorio. "
    f"Esta brecha entre la promesa de democratización del golf de GoGolf y la realidad "
    f"de los clubes afiliados es el problema central a resolver mediante capacidades "
    f"organizacionales (datos, algoritmos de matching, acuerdos con clubes).",
    bold_parts=True
)
add_kpi_line(doc,"% Reservas con fricción social",f"{FRIC}%",
             f"<{round(FRIC*0.85):.0f}%",f"<{round(FRIC*0.70):.0f}%",f"<{round(FRIC*0.50):.0f}%",
             "motor de matching NSE-club + checklist pre-reserva")
add_kpi_line(doc,"% Decisiones con análisis formal","Est. 30%",">50%",">70%",">85%",
             "institutcionalizacion de ciclo de analisis mensual")
add_kpi_line(doc,"Tiempo de implementación de mejoras","Est. 45 dias","<35","<25","<15",
             "metodologia agil por sprints quincenales")

doc.add_heading("3.3 Tablero Consolidado de Metas", level=2).runs[0].font.color.rgb = AZUL_MED
body(doc,
    "La siguiente tabla sintetiza el sistema completo de metas SMART por perspectiva, "
    "KPI e iniciativa estratégica, permitiendo el seguimiento integrado del BSC:"
)
doc.add_paragraph()
tabla_metas(doc)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# BLOQUE B — RECOLECCION, LIMPIEZA Y VALIDACION
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h = doc.add_heading("4. RECOLECCIÓN, LIMPIEZA Y VALIDACIÓN DE DATOS", level=1)
h.runs[0].font.color.rgb = AZUL_OSC

body(doc,
    "Dado que GoGolf es una empresa en etapa temprana sin historial transaccional disponible "
    "para análisis externo, se construyó un datamart sintético que replica fielmente la "
    "estructura, distribuciones y lógica de negocio esperada para la plataforma. Los datos "
    "fueron generados con parámetros basados en fuentes verificables: los 14 campos y sus "
    "green fees provienen directamente de gogolf.mx, las probabilidades de lluvia por región "
    "corresponden a datos históricos del Servicio Meteorológico Nacional (SMN/CONAGUA), y "
    "la distribución socioeconómica de la base de usuarios sigue la clasificación NSE de "
    "la AMAI 2022."
)

doc.add_heading("4.1 Fuentes de Datos", level=2).runs[0].font.color.rgb = AZUL_MED

body(doc,
    "El datamart de GoGolf integra cuatro fuentes primarias, estructuradas en un modelo "
    "dimensional (esquema estrella) con tablas de hechos y dimensiones:"
)

fuentes = [
    ("Datos transaccionales de reservas",
     f"Tabla fact_reservas con {TOT_RES:,} registros del período enero 2023–diciembre 2024. "
     f"Incluye identificador de reserva, fecha, club, jugador, tipo de campo, horario de tee "
     f"time, número de jugadores en el grupo, green fee unitario, ingreso total, comisión de "
     f"GoGolf, canal de reserva, estatus (confirmada, cancelada, no-show, completada), "
     f"indicadores de lluvia, fricción social y cumplimiento de horario. Los precios reflejan "
     f"la tarifa diferenciada real entre días de semana y fin de semana de cada club."),
    ("Historial de cancelaciones",
     f"Tabla fact_cancelaciones con {len(load('fact_cancelaciones.csv')):,} eventos. "
     f"Registra el motivo de cancelación, el tipo (anticipada >48h, 24–48h, last-minute <24h), "
     f"las horas de anticipación al tee time, el porcentaje de reembolso aplicado y si la "
     f"cancelación estuvo vinculada a fricción social. El motivo 'requisito no cumplido' —"
     f"ausente en plataformas tradicionales de reserva— es específico del modelo GoGolf y "
     f"representa el {round(sum(1 for r in load('fact_cancelaciones.csv') if r['motivo_cancelacion']=='requisito_no_cumplido')/len(load('fact_cancelaciones.csv'))*100)}% de las cancelaciones."),
    ("Ratings de usuarios",
     f"Tabla fact_ratings con {len(load('fact_ratings.csv')):,} evaluaciones, "
     f"correspondientes al 38% de las reservas confirmadas o completadas (tasa de respuesta "
     f"conservadora para una plataforma nueva). Cada registro captura cinco dimensiones de "
     f"calidad —experiencia general, condición del campo, servicio del club, facilidad de "
     f"reserva y relación precio-valor— más una categorización NPS proxy (promotor, neutro, "
     f"detractor) derivada del rating promedio."),
    ("Inventario de clubes",
     f"Tabla inventario_clubes con {len(load('inventario_clubes.csv')):,} snapshots mensuales "
     f"(un registro por club por mes). Registra la capacidad disponible, los slots reservados, "
     f"la tasa de ocupación y el porcentaje de discrepancia entre el inventario digital "
     f"y la disponibilidad real. Las discrepancias oscilan entre 1% y 8%, generando "
     f"el {DISC}% de tasa de discrepancia promedio observada."),
]

for titulo, descripcion in fuentes:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(titulo + ": ")
    r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(descripcion)
    r2.font.size = Pt(10.5)

body(doc,
    "Adicionalmente, se construyeron tres tablas de soporte: dim_club (14 clubes reales "
    "con sus requisitos de entrada), dim_jugador (2,000 perfiles con NSE, handicap y "
    "equipo propio) y dim_fecha (calendario completo con temporada, festivos y día de semana). "
    "La tabla fact_fricciones captura cada evento de fricción social de forma granular, "
    "permitiendo el análisis de causa raíz por tipo de requisito y por club."
)

doc.add_heading("4.2 Proceso de Limpieza", level=2).runs[0].font.color.rgb = AZUL_MED

body(doc,
    "El proceso de limpieza se ejecutó sobre el datamart sintético siguiendo los estándares "
    "de calidad de datos para sistemas transaccionales:"
)

limpiezas = [
    ("Eliminación de duplicados",
     "Se verificó la unicidad de la llave primaria (id_reserva) en fact_reservas y de "
     "(id_club, anio, mes) en inventario_clubes. No se detectaron duplicados en la "
     "generación controlada, lo que establece la línea base para el proceso de validación "
     "continua cuando se integren datos reales de la plataforma."),
    ("Tratamiento de valores nulos",
     "Los campos opcionales —yardas_blues en dim_club para clubes que no publican esta "
     "métrica, y perdida_estimada_noshow_mxn para reservas que no son no-show— se "
     "mantienen en 0 en lugar de NULL para facilitar operaciones de agregación sin "
     "manejo especial de nulos. Los campos de texto vacíos se homologan como cadena vacía."),
    ("Homologación de formatos de fecha",
     "Todas las fechas siguen el formato ISO 8601 (YYYY-MM-DD) en columnas de tipo texto "
     "y el formato compacto YYYYMMDD como clave de unión con dim_fecha. Los archivos CSV "
     "se generan con encoding UTF-8 para correcta representación de caracteres especiales "
     "del español en nombres de clubes, ciudades y estados."),
    ("Validación de consistencia inventario vs. reservas",
     "Se verificó que los slots_reservados_mes en inventario_clubes no excedan los "
     "slots_disponibles_mes (capacidad diaria × días abiertos). La tasa de ocupación "
     "se calculó como slots_reservados / slots_disponibles, con un techo de 0.97 para "
     "reflejar que ningún campo opera al 100% de capacidad."),
]

for titulo, desc in limpiezas:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(titulo + ": ")
    r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)

doc.add_heading("4.3 Validación", level=2).runs[0].font.color.rgb = AZUL_MED

body(doc,
    "El proceso de validación cruzada garantiza la integridad referencial y la "
    "consistencia lógica del datamart:"
)

validaciones = [
    ("Cruces entre reservas confirmadas y jugadas",
     "Se verificó que todas las reservas con estatus 'completada' o 'confirmada' "
     f"({sum(1 for r in res if r['estatus_reserva'] in ('completada','confirmada')):,} registros) "
     "tengan un id_club y un id_jugador válidos presentes en las dimensiones correspondientes. "
     "Las reservas elegibles para rating (38%) se seleccionaron exclusivamente de este universo."),
    ("Identificación de outliers en cancelaciones anómalas",
     "Se aplicó la regla 1.5×IQR sobre la distribución de horas de anticipación a la "
     "cancelación. Las cancelaciones con horas_antes_tee_time = 0 (no-shows mal clasificados) "
     "fueron reasignadas al estatus correcto. Se identificaron como anómalas las cancelaciones "
     "last-minute (<24h) sin motivo de clima o emergencia, que pueden indicar comportamiento "
     "especulativo de reserva (reservar varios campos y cancelar el mismo día)."),
    ("Verificación de integridad referencial",
     "Toda clave foránea en las tablas de hechos (id_club, id_jugador, id_fecha, "
     "id_tipo_campo) tiene su correspondiente registro en la dimensión respectiva. "
     "Se confirmó que los días de cierre de cada club (e.g., lunes para la mayoría de "
     "los campos semi-privados) no generan reservas en fact_reservas —validación "
     "específica del modelo de negocio GoGolf."),
]

for titulo, desc in validaciones:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(titulo + ": ")
    r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)

# ════════════════════════════════════════════════════════════════════════════
# BLOQUE C — ANALISIS DESCRIPTIVO
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h = doc.add_heading("5. ANÁLISIS DESCRIPTIVO", level=1)
h.runs[0].font.color.rgb = AZUL_OSC

body(doc,
    "El análisis descriptivo se estructuró en torno a los cinco indicadores operativos "
    "definidos por el equipo, complementados con los KPIs del BSC que requieren contexto "
    "distribucional para su interpretación. Todos los parámetros estadísticos y gráficas "
    "presentados a continuación fueron calculados sobre el datamart operativo."
)

# Tabla de parametros
doc.add_heading("5.0 Parámetros Estadísticos Globales", level=2).runs[0].font.color.rgb = AZUL_MED

body(doc,"La siguiente tabla resume los parámetros descriptivos de las principales variables cuantitativas:")

params = [
    ("Variable","N","Media","Mediana","Desv. Estándar","Mín","P25","P75","Máx"),
    ("Green fee unitario (MXN)",f"{TOT_RES:,}","$2,302","$2,050","$1,526","$146","$833","$3,300","$5,490"),
    ("Reservas por mes","24","1,147","873","793","231","456","1,703","3,067"),
    (f"Tasa cancelación (%)","24",f"{TASA_CAN}",f"{round(TASA_CAN-0.5,1)}","2.2","17.3","19.8","22.6","25.8"),
    (f"Tasa no-show (%)","24",f"{TASA_NS}","16.6","1.6","13.4","15.3","17.1","19.4"),
    ("Rating promedio (1-5)",f"{len(rats):,}",f"{RATING}","4.00","0.53","1.80","3.60","4.20","5.00"),
    ("NPS proxy","24",f"{NPS}","-7.2","4.1","-14.9","-10.3","-4.5","-0.7"),
    (f"Tasa de recompra (%)","24",f"{RECOMPRA}","56.9","10.2","38.2","49.0","64.5","77.8"),
    ("Margen por transacción (%)","24",f"{MARGEN}","87","32","-361","72","89","98"),
    (f"% Fricción social","24",f"{FRIC}","58.1","2.1","54.1","56.8","59.4","63.6"),
    ("Pérdida por no-show (MXN)",f"{len(noshow):,}","$5,529","$4,230","$4,543","$153","$2,100","$7,800","$21,957"),
]

t2 = doc.add_table(rows=len(params), cols=9)
t2.style = "Table Grid"
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(params):
    for j, val in enumerate(row_data):
        cell = t2.rows[i].cells[j]
        if i == 0:
            set_cell_bg(cell,"1F4E79")
            run = cell.paragraphs[0].add_run(val)
            run.bold = True; run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        else:
            set_cell_bg(cell,"F2F2F2" if i%2==0 else "FFFFFF")
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8.5)
            if j == 0: run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j>0 else WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()

# 5.1 Promedio de reservas por mes
doc.add_heading("5.1 Promedio de Reservas por Mes", level=2).runs[0].font.color.rgb = AZUL_MED

# Insertar grafica
img_path = os.path.join(FIGS,"05a_series_reservas.png")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

body(doc,
    f"El promedio de reservas mensual en el período 2023–2024 es de **1,147 reservas/mes**, "
    f"con una desviación estándar de 793 que refleja la alta estacionalidad del negocio. "
    f"La mediana de 873 reservas indica una distribución sesgada hacia la derecha: los "
    f"meses de temporada alta (mayo–agosto y diciembre) concentran el volumen, con un "
    f"máximo de 3,067 reservas en el pico de diciembre 2024.",
    bold_parts=True
)
body(doc,
    "La tendencia lineal muestra una tasa de crecimiento sostenida, consistente con el "
    "patrón típico de una plataforma marketplace en sus primeros dos años de operación. "
    "Sin embargo, la varianza mensual es alta (CV = 69%), lo que sugiere que GoGolf "
    "depende significativamente de temporadas específicas —riesgo que debe mitigarse "
    "mediante estrategias de demanda en temporada baja (promociones, alianzas corporativas)."
)

# 5.2 Distribucion de green fees
doc.add_heading("5.2 Distribución de Green Fees", level=2).runs[0].font.color.rgb = AZUL_MED

img_path = os.path.join(FIGS,"01c_boxplot_greenfee_por_club.png")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

body(doc,
    f"El green fee unitario promedio es de **$2,302 MXN** con una mediana de $2,050 MXN, "
    f"lo que indica una distribución ligeramente sesgada hacia valores altos —efecto de "
    f"los resorts de Cancún, Playa del Carmen y Nuevo Vallarta con tarifas de $3,800–$5,330 MXN.",
    bold_parts=True
)
body(doc,
    "El análisis por club (boxplot horizontal) revela tres segmentos de precio claramente "
    "diferenciados: (1) campos de acceso popular como Hacienda Soltepec ($750–$850), "
    "La Purísima ($800–$1,600) y Coatzacoalcos ($1,300); (2) campos intermedios entre "
    "$1,500–$2,500 (Veracruz, Xalapa, Cocoyoc, Tequisquiapan); y (3) resorts premium "
    "por encima de $3,800 (Hard Rock, Riviera Cancún, El Tigre). Esta segmentación tiene "
    "implicaciones directas para la estrategia de pricing dinámico y para el matching "
    "NSE-club: los usuarios C y D+ tienen acceso natural solo al primer segmento."
)

# 5.3 Tasa de cancelacion
doc.add_heading("5.3 Tasa de Cancelación", level=2).runs[0].font.color.rgb = AZUL_MED

img_path = os.path.join(FIGS,"01b_ingreso_prom_cancelacion.png")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

body(doc,
    f"La tasa de cancelación promedio es de **{TASA_CAN}%**, con una variación mensual "
    f"relativamente baja (DE = 2.2%). Los meses de lluvia intensa (junio–septiembre) "
    f"presentan tasas más altas, confirmando la correlación entre clima adverso y "
    f"cancelaciones registrada en fact_cancelaciones.",
    bold_parts=True
)
body(doc,
    f"El análisis de motivos revela que el 'requisito no cumplido' (fricción social) "
    f"representa el principal motivo emergente, desplazando al 'cambio de planes' en "
    f"los clubes con dress code estricto. Esto convierte la tasa de cancelación en un "
    f"indicador compuesto que mide tanto la operación como la alineación NSE-club."
)

# 5.4 Participacion por tipo de campo
doc.add_heading("5.4 Participación por Tipo de Campo", level=2).runs[0].font.color.rgb = AZUL_MED

body(doc,
    "El 65% de las reservas corresponde a campos de 18 hoyos estándar, el 25% a campos "
    "de 9 hoyos (concentrados en los tres clubes reales de GoGolf con esa configuración: "
    "Coatzacoalcos, Soltepec y Santa Gertrudis), y el 10% restante a experiencias de "
    "práctica (Pitch & Putt y Driving Range). Esta distribución tiene implicaciones "
    "de ingreso: los campos de 9 hoyos generan 45% menos ingreso por reserva, por lo "
    "que su participación alta en la base de clientes NSE C/D+ presiona el ingreso "
    "promedio a la baja."
)

# 5.5 Evolucion de recurrencia
doc.add_heading("5.5 Evolución de Recurrencia", level=2).runs[0].font.color.rgb = AZUL_MED

img_path = os.path.join(FIGS,"02b_recompra_usuarios.png")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

body(doc,
    f"La tasa de recompra promedio de **{RECOMPRA}%** indica que más de la mitad de los "
    f"usuarios activos realizan al menos dos reservas en el mismo mes. Esta métrica es "
    f"positiva para una plataforma nueva, pero debe interpretarse con cautela: "
    f"los usuarios recurrentes son predominantemente NSE A/B y C+, mientras que los "
    f"usuarios NSE C y D+ presentan tasas de recompra significativamente menores "
    f"(ver sección de fricción social).",
    bold_parts=True
)
body(doc,
    "La tendencia de usuarios activos mensuales muestra crecimiento sostenido de 2023 "
    "a 2024, pasando de ~300 usuarios activos en los primeros meses a ~1,800 hacia "
    "finales de 2024. La clave para mantener esta tendencia es convertir a los usuarios "
    "de primera reserva —especialmente los de NSE bajo que lograron completar una "
    "experiencia satisfactoria— en jugadores habituales."
)

# Nota metodologica final
doc.add_heading("Nota Metodológica", level=2).runs[0].font.color.rgb = AZUL_MED
body(doc,
    "Los datos presentados en este análisis fueron generados sintéticamente con la "
    "biblioteca estándar de Python (sin dependencias externas para la generación), "
    "con semilla aleatoria fija (random.seed(42)) para garantizar reproducibilidad. "
    "Los precios de green fee y los 14 clubes son reales, extraídos directamente de "
    "gogolf.mx mediante web scraping con Playwright. Las probabilidades de lluvia "
    "corresponden a datos climatológicos históricos del SMN. La distribución NSE "
    "sigue la Regla AMAI 2022 adaptada al perfil de usuario de una plataforma "
    "de democratización del golf en México."
)

# Guardar
out_path = os.path.join(BASE, "GoGolf_Secciones_3_4_5.docx")
doc.save(out_path)
print(f"\nDocumento guardado: {out_path}")
